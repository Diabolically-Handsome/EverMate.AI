"""Tests for engine.storage: file helpers, MemoryStore, InstanceLock."""

from __future__ import annotations

import os

import pytest

from engine.storage import (
    VAULT_HEADER,
    InstanceLock,
    MemoryStore,
    append_text,
    read_text,
    safe_filename,
    write_text,
)


# ---------------- file helpers ----------------


class TestFileHelpers:
    def test_read_text_missing_file_returns_empty(self, tmp_path):
        assert read_text(str(tmp_path / "nope.txt")) == ""

    def test_write_text_roundtrip(self, tmp_path):
        p = str(tmp_path / "sub" / "a.txt")  # parent dir auto-created
        write_text(p, "你好 hello\n")
        assert read_text(p) == "你好 hello\n"

    def test_write_text_overwrites(self, tmp_path):
        p = str(tmp_path / "a.txt")
        write_text(p, "first")
        write_text(p, "second")
        assert read_text(p) == "second"

    def test_write_text_leaves_no_temp_files(self, tmp_path):
        p = str(tmp_path / "a.txt")
        write_text(p, "content")
        leftovers = [f for f in os.listdir(tmp_path) if f.startswith(".evermate-")]
        assert leftovers == []

    def test_append_text(self, tmp_path):
        p = str(tmp_path / "a.txt")
        append_text(p, "one\n")
        append_text(p, "two\n")
        assert read_text(p) == "one\ntwo\n"


class TestSafeFilename:
    def test_slashes_replaced(self):
        assert safe_filename("my/file\\name.txt") == "my_file_name.txt"

    def test_weird_chars_replaced(self):
        assert safe_filename("bad|name?.txt") == "bad_name_.txt"

    def test_chinese_and_common_chars_preserved(self):
        assert safe_filename("重要笔记 (2).txt") == "重要笔记 (2).txt"

    def test_long_names_truncated(self):
        assert len(safe_filename("a" * 300 + ".txt")) == 180


# ---------------- MemoryStore: schema ----------------


class TestSchema:
    def test_schema_version_is_2(self, store):
        row = store.conn.execute(
            "SELECT value FROM meta WHERE key='schema_version'"
        ).fetchone()
        assert row["value"] == "2"

    def test_expected_tables_exist(self, store):
        names = {
            r["name"]
            for r in store.conn.execute(
                "SELECT name FROM sqlite_master WHERE type='table'"
            ).fetchall()
        }
        assert {"chunks", "terms", "postings", "meta", "imported_files"} <= names

    def test_wal_mode(self, store):
        row = store.conn.execute("PRAGMA journal_mode;").fetchone()
        assert str(row[0]).lower() == "wal"

    def test_migration_drops_redundant_term_index(self, tmp_path):
        d = str(tmp_path / "m")
        s = MemoryStore(d)
        # Simulate a v1 database that still carries idx_postings_term.
        s.conn.execute(
            "CREATE INDEX IF NOT EXISTS idx_postings_term ON postings(term)"
        )
        s.conn.execute("UPDATE meta SET value='1' WHERE key='schema_version'")
        s.conn.commit()
        s.close()

        s2 = MemoryStore(d)
        try:
            row = s2.conn.execute(
                "SELECT name FROM sqlite_master WHERE type='index' AND name='idx_postings_term'"
            ).fetchone()
            assert row is None
            ver = s2.conn.execute(
                "SELECT value FROM meta WHERE key='schema_version'"
            ).fetchone()
            assert ver["value"] == "2"
        finally:
            s2.close()


# ---------------- MemoryStore: chunks & stats ----------------


class TestAddChunk:
    def test_add_chunk_persists_text_and_index(self, store):
        cid = store.add_chunk("the quick brown fox", source="chat")
        assert cid >= 1
        assert store.count_chunks() == 1
        assert store.chunk_text_by_id(cid) == "the quick brown fox"
        # chunk file on disk
        row = store.chunk_row_by_id(cid)
        assert os.path.exists(os.path.join(store.memory_dir, str(row["path"])))
        # vault summary appended
        assert f"Chunk {cid:08d}" in read_text(store.vault_md_path)

    def test_empty_text_rejected(self, store):
        assert store.add_chunk("   ", source="chat") == -1

    def test_stopword_only_text_rejected(self, store):
        # tokenizes to nothing -> not indexable
        assert store.add_chunk("的 了", source="chat") == -1
        assert store.count_chunks() == 0

    def test_corpus_stats_cached_and_invalidated(self, store):
        store.add_chunk("alpha beta gamma", source="chat")
        n1, avgdl1 = store.corpus_stats()
        assert n1 == 1
        assert avgdl1 > 0
        # cache returns the same values without a write
        assert store.corpus_stats() == (n1, avgdl1)
        store.add_chunk("delta epsilon", source="chat")
        n2, _ = store.corpus_stats()
        assert n2 == 2


# ---------------- MemoryStore: ingestion ----------------


class TestIngestion:
    def test_ingest_txt_in_multiple_chunks(self, tmp_path):
        s = MemoryStore(str(tmp_path / "m"), chunk_chars=80)
        try:
            src = tmp_path / "doc.txt"
            src.write_text(
                "\n".join(f"line number {i} with some words" for i in range(12)) + "\n",
                encoding="utf-8",
            )
            made = s.ingest_file(str(src), source="upload:doc.txt")
            assert made >= 2
            assert s.count_chunks() == made
        finally:
            s.close()

    def test_ingest_unknown_extension_is_noop(self, store, tmp_path):
        src = tmp_path / "doc.pdf"
        src.write_bytes(b"%PDF-1.4 fake")
        assert store.ingest_file(str(src), source="upload:doc.pdf") == 0
        assert store.count_chunks() == 0

    def test_ingest_docx(self, tmp_path):
        pytest.importorskip("docx")
        from docx import Document

        doc = Document()
        doc.add_paragraph("我们在2024年4月3日去了京都。")
        doc.add_paragraph("The cherry blossoms were lovely.")
        src = tmp_path / "trip.docx"
        doc.save(str(src))

        s = MemoryStore(str(tmp_path / "m"))
        try:
            made = s.ingest_file(str(src), source="upload:trip.docx")
            assert made == 1
            text = s.chunk_text_by_id(1)
            assert "京都" in text
            assert "cherry blossoms" in text
        finally:
            s.close()

    def test_ingest_text_chunks_long_input(self, tmp_path):
        s = MemoryStore(str(tmp_path / "m"), chunk_chars=100)
        try:
            text = "\n".join(f"sentence {i} about nothing much" for i in range(20))
            made = s.ingest_text(text, source="chat")
            assert made >= 2
            assert s.count_chunks() == made
        finally:
            s.close()

    def test_ingest_empty_text(self, store):
        assert store.ingest_text("", source="chat") == 0


# ---------------- MemoryStore: uploads & dedup ----------------


class TestUploads:
    def test_import_copies_into_uploads(self, store, tmp_path):
        src = tmp_path / "notes.txt"
        src.write_text("hello memory", encoding="utf-8")
        stored = store.import_files([str(src)])
        assert len(stored) == 1
        assert os.path.dirname(stored[0]) == store.uploads_dir
        assert read_text(stored[0]) == "hello memory"

    def test_reimporting_identical_content_returns_empty(self, store, tmp_path):
        src1 = tmp_path / "a.txt"
        src1.write_text("identical content", encoding="utf-8")
        assert len(store.import_files([str(src1)])) == 1

        # same bytes under a different name -> deduped by sha256
        src2 = tmp_path / "b.txt"
        src2.write_text("identical content", encoding="utf-8")
        assert store.import_files([str(src2)]) == []

        # different content is accepted
        src3 = tmp_path / "c.txt"
        src3.write_text("different content", encoding="utf-8")
        assert len(store.import_files([str(src3)])) == 1

    def test_unsupported_and_missing_files_skipped(self, store, tmp_path):
        pdf = tmp_path / "x.pdf"
        pdf.write_bytes(b"fake")
        missing = tmp_path / "does-not-exist.txt"
        assert store.import_files([str(pdf), str(missing), ""]) == []

    def test_mark_ingested(self, store, tmp_path):
        src = tmp_path / "a.txt"
        src.write_text("content to ingest", encoding="utf-8")
        stored = store.import_files([str(src)])[0]
        row = store.conn.execute("SELECT ingested FROM imported_files").fetchone()
        assert row["ingested"] == 0
        store.mark_ingested(stored)
        row = store.conn.execute("SELECT ingested FROM imported_files").fetchone()
        assert row["ingested"] == 1

    def test_list_uploads(self, store, tmp_path):
        assert store.list_uploads() == []
        src = tmp_path / "a.txt"
        src.write_text("text", encoding="utf-8")
        stored = store.import_files([str(src)])[0]
        assert store.list_uploads() == [stored]

    def test_delete_upload_allows_reimport(self, store, tmp_path):
        src = tmp_path / "a.txt"
        src.write_text("delete me", encoding="utf-8")
        stored = store.import_files([str(src)])[0]
        assert store.delete_upload(stored) is True
        assert not os.path.exists(stored)
        # dedup record went away too, so the same content imports again
        assert len(store.import_files([str(src)])) == 1

    def test_delete_upload_refuses_paths_outside_uploads(self, store, tmp_path):
        outside = tmp_path / "outside.txt"
        outside.write_text("not yours", encoding="utf-8")
        assert store.delete_upload(str(outside)) is False
        assert os.path.exists(outside)


# ---------------- MemoryStore: destructive ops ----------------


class TestDestructiveOps:
    def test_reset_index_backs_up_db_and_clears(self, store):
        store.add_chunk("some indexed text here", source="chat")
        assert store.count_chunks() == 1
        store.reset_index()
        assert store.count_chunks() == 0
        assert os.path.exists(store.db_path + ".bak")
        assert os.listdir(store.chunks_dir) == []
        assert read_text(store.vault_md_path) == VAULT_HEADER

    def test_clear_chat_history_empties_logs(self, store):
        append_text(store.chat_log_path, "[ts] user: hi\n")
        append_text(store.buffer_path, "user: hi\n")
        store.clear_chat_history()
        assert read_text(store.chat_log_path) == ""
        assert read_text(store.buffer_path) == ""

    def test_wipe_all_removes_everything_and_stays_usable(self, store, tmp_path):
        store.add_chunk("remember this fact", source="chat")
        src = tmp_path / "a.txt"
        src.write_text("uploaded doc", encoding="utf-8")
        store.import_files([str(src)])

        store.wipe_all()

        assert store.count_chunks() == 0
        assert store.count_terms() == 0
        assert store.list_uploads() == []
        assert read_text(store.vault_md_path) == VAULT_HEADER
        assert not os.path.exists(store.db_path + ".bak")
        # the store remains usable after the wipe
        assert store.add_chunk("fresh start", source="chat") >= 1


# ---------------- InstanceLock ----------------


class TestInstanceLock:
    def test_second_lock_on_same_dir_fails(self, tmp_path):
        # Two InstanceLock objects open separate fds; flock makes them
        # conflict even within a single process.
        lock1 = InstanceLock(str(tmp_path))
        lock2 = InstanceLock(str(tmp_path))
        try:
            assert lock1.acquire() is True
            assert lock2.acquire() is False
        finally:
            lock1.release()

        # released -> the other instance can take over
        try:
            assert lock2.acquire() is True
        finally:
            lock2.release()

    def test_release_without_acquire_is_safe(self, tmp_path):
        InstanceLock(str(tmp_path)).release()  # must not raise
